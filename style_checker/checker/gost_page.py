"""
checker/gost_page.py
Проверка требований ГОСТ 7.32-2017 / ГОСТ Р 7.0.11-2011 к оформлению страницы:
  - Поля страницы
  - Шрифт и кегль основного текста
  - Межстрочный интервал
  - Абзацный отступ (красная строка)

Все значения читаются напрямую из XML внутри DOCX — без сторонних библиотек,
кроме python-docx.

Единицы измерения в DOCX XML:
  - Поля:   twips (1 twip = 1/1440 дюйма = 0.01763 мм)
  - Отступ: twips (w:firstLine)
  - Кегль:  half-points (w:sz, делить на 2 → pt)
  - Интервал: w:line / 240 при lineRule="auto" = множитель (1.0, 1.5, 2.0…)
"""

from .base import BaseChecker
from ..models.report import CheckCategory, Severity, CheckResult
from ..core.document import ParsedDocument
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from io import BytesIO


# ── Константы ГОСТ ──────────────────────────────────────────────────────────

MM_PER_TWIP = 25.4 / 1440   # 1 twip → мм
PT_PER_HALFPT = 0.5         # half-point → pt


def _twips_to_mm(value) -> float | None:
    """Конвертирует строку или число (twips) в мм."""
    try:
        return round(int(float(value)) * MM_PER_TWIP, 1)
    except (TypeError, ValueError):
        return None


def _halfpt_to_pt(value) -> float | None:
    """Конвертирует half-points в pt (w:sz → pt)."""
    try:
        return int(value) * PT_PER_HALFPT
    except (TypeError, ValueError):
        return None


# ── Вспомогательные функции чтения XML ──────────────────────────────────────

def _get_page_margins(doc: DocxDocument) -> dict | None:
    """Читает поля страницы из первой секции."""
    for section in doc.sections:
        pgMar = section._sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            return {
                "left":   _twips_to_mm(pgMar.get(qn("w:left"))),
                "right":  _twips_to_mm(pgMar.get(qn("w:right"))),
                "top":    _twips_to_mm(pgMar.get(qn("w:top"))),
                "bottom": _twips_to_mm(pgMar.get(qn("w:bottom"))),
            }
    return None


def _get_doc_default_font(doc: DocxDocument) -> dict:
    """
    Читает шрифт и кегль из docDefaults — они применяются ко всему документу,
    если в параграфе/run нет явного переопределения.
    """
    result = {"name": None, "size_pt": None}
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn("w:docDefaults"))
    if docDefaults is None:
        return result

    rPrDefault = docDefaults.find(".//" + qn("w:rPrDefault"))
    if rPrDefault is None:
        return result

    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        return result

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        result["name"] = (
            rFonts.get(qn("w:ascii"))
            or rFonts.get(qn("w:hAnsi"))
            or rFonts.get(qn("w:cs"))
        )

    sz = rPr.find(qn("w:sz"))
    if sz is not None:
        result["size_pt"] = _halfpt_to_pt(sz.get(qn("w:val")))

    return result


def _get_para_spacing(para_element) -> dict:
    """
    Извлекает межстрочный интервал параграфа.
    Возвращает: {"line": int|None, "rule": str|None, "multiplier": float|None}
    """
    pPr = para_element.find(qn("w:pPr"))
    if pPr is None:
        return {"line": None, "rule": None, "multiplier": None}

    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        return {"line": None, "rule": None, "multiplier": None}

    line = sp.get(qn("w:line"))
    rule = sp.get(qn("w:lineRule"))

    multiplier = None
    if line and rule == "auto":
        # 240 = single, 360 = 1.5, 480 = double
        multiplier = round(int(line) / 240, 2)

    return {"line": line, "rule": rule, "multiplier": multiplier}


def _get_para_first_line_indent(para_element) -> float | None:
    """
    Читает отступ первой строки (красная строка) в мм.
    w:ind/@w:firstLine — в twips.
    """
    pPr = para_element.find(qn("w:pPr"))
    if pPr is None:
        return None

    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return None

    val = ind.get(qn("w:firstLine"))
    return _twips_to_mm(val)


def _get_para_explicit_font(para_element) -> dict:
    """
    Возвращает явно заданный шрифт/кегль параграфа (из w:pPr/w:rPr или первого run).
    Если явно не задан — возвращает None (наследуется от docDefaults).
    """
    result = {"name": None, "size_pt": None}

    # Сначала проверяем rPr внутри pPr (стиль параграфа)
    pPr = para_element.find(qn("w:pPr"))
    if pPr is not None:
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                result["name"] = (
                    rFonts.get(qn("w:ascii"))
                    or rFonts.get(qn("w:hAnsi"))
                    or rFonts.get(qn("w:cs"))
                )
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                result["size_pt"] = _halfpt_to_pt(sz.get(qn("w:val")))

    # Затем — из runs
    for run in para_element.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is not None:
            if result["name"] is None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    result["name"] = (
                        rFonts.get(qn("w:ascii"))
                        or rFonts.get(qn("w:hAnsi"))
                        or rFonts.get(qn("w:cs"))
                    )
            if result["size_pt"] is None:
                sz = rPr.find(qn("w:sz"))
                if sz is not None:
                    result["size_pt"] = _halfpt_to_pt(sz.get(qn("w:val")))
        if result["name"] and result["size_pt"]:
            break

    return result


# ── Сам чекер ────────────────────────────────────────────────────────────────

class GostPageChecker(BaseChecker):
    """
    Проверяет соответствие оформления страницы требованиям ГОСТ 7.32-2017.
    Для работы требует доступ к объекту Document из python-docx,
    поэтому принимает его через конструктор или извлекает из ParsedDocument.
    """

    def __init__(self, config: dict, docx_source=None):
        """
        Args:
            config:      конфиг правил (секция "gost_page" из rules.json)
            docx_source: путь к файлу / bytes / BytesIO.
                         Если None — парсер попробует использовать последний
                         открытый файл (для совместимости с движком).
        """
        super().__init__(config)
        self._docx_source = docx_source

    @property
    def name(self) -> str:
        return "Проверка оформления страницы (ГОСТ)"

    @property
    def category(self) -> CheckCategory:
        return CheckCategory.FORMATTING

    def check(self, document: ParsedDocument) -> CheckResult:
        result = self._result()

        if self._docx_source is None:
            result.violations.append(self._violation(
                rule_id="GOST-NO-SOURCE",
                severity=Severity.ERROR,
                message="Чекер ГОСТ не получил источник DOCX-файла",
            ))
            return result

        src = self._docx_source
        if isinstance(src, bytes):
            src = BytesIO(src)
        doc = DocxDocument(src)

        cfg = self.config

        # ── 1. Поля страницы ────────────────────────────────────────
        margins = _get_page_margins(doc)
        if margins:
            self._check_margin(result, margins, "left",   cfg.get("margin_left_mm",   30.0), "левое")
            self._check_margin(result, margins, "right",  cfg.get("margin_right_mm",  10.0), "правое")
            self._check_margin(result, margins, "top",    cfg.get("margin_top_mm",    20.0), "верхнее")
            self._check_margin(result, margins, "bottom", cfg.get("margin_bottom_mm", 20.0), "нижнее")
        else:
            result.violations.append(self._violation(
                rule_id="GOST-MARGINS-NOT-FOUND",
                severity=Severity.WARNING,
                message="Не удалось прочитать поля страницы из документа",
            ))

        # ── 2. Шрифт и кегль ────────────────────────────────────────
        default_font = _get_doc_default_font(doc)
        required_font      = cfg.get("font_name", "Times New Roman")
        required_size_pt   = cfg.get("font_size_pt", 14.0)
        tolerance_pt       = cfg.get("font_size_tolerance_pt", 0.5)
        allowed_fonts      = cfg.get("allowed_fonts", [required_font])

        # Проверяем дефолтный шрифт документа
        if default_font["name"]:
            if default_font["name"] not in allowed_fonts:
                result.violations.append(self._violation(
                    rule_id="GOST-FONT-NAME",
                    severity=Severity.WARNING,
                    message=(
                        f"Шрифт документа по умолчанию: «{default_font['name']}» — "
                        f"ожидается «{required_font}»"
                    ),
                    suggestion=f"Установите шрифт «{required_font}» для всего документа",
                ))
            else:
                result.passed += 1

        if default_font["size_pt"]:
            if abs(default_font["size_pt"] - required_size_pt) > tolerance_pt:
                result.violations.append(self._violation(
                    rule_id="GOST-FONT-SIZE-DEFAULT",
                    severity=Severity.WARNING,
                    message=(
                        f"Размер шрифта по умолчанию: {default_font['size_pt']} pt — "
                        f"ожидается {required_size_pt} pt"
                    ),
                    suggestion=f"Установите кегль {required_size_pt} pt для всего документа",
                ))
            else:
                result.passed += 1

        # Проверяем явно заданные шрифты в параграфах основного текста
        font_violations_names = []
        font_violations_sizes = []
        checked_paras = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            # Пропускаем заголовки — у них другой кегль по стандарту
            if "heading" in style_name.lower() or "заголовок" in style_name.lower():
                continue
            checked_paras += 1

            explicit = _get_para_explicit_font(para._element)

            # Шрифт
            font = explicit["name"] or default_font["name"]
            if font and font not in allowed_fonts:
                font_violations_names.append((para, font))

            # Кегль
            size = explicit["size_pt"] or default_font["size_pt"]
            if size and abs(size - required_size_pt) > tolerance_pt:
                font_violations_sizes.append((para, size))

        # Сводим нарушения шрифта — не спамим на каждый параграф
        if font_violations_names:
            unique_fonts = {f for _, f in font_violations_names}
            result.violations.append(self._violation(
                rule_id="GOST-FONT-NAME-BODY",
                severity=Severity.WARNING,
                message=(
                    f"Основной текст использует недопустимые шрифты: "
                    f"{', '.join(f'«{f}»' for f in unique_fonts)} "
                    f"(затронуто параграфов: {len(font_violations_names)})"
                ),
                suggestion=f"Используйте только «{required_font}» в основном тексте",
                context=font_violations_names[0][0].text[:80],
            ))
        elif checked_paras > 0:
            result.passed += 1

        if font_violations_sizes:
            unique_sizes = {s for _, s in font_violations_sizes}
            result.violations.append(self._violation(
                rule_id="GOST-FONT-SIZE-BODY",
                severity=Severity.WARNING,
                message=(
                    f"Некорректный размер шрифта в основном тексте: "
                    f"{', '.join(str(s) + ' pt' for s in unique_sizes)} "
                    f"(затронуто параграфов: {len(font_violations_sizes)})"
                ),
                suggestion=f"Используйте кегль {required_size_pt} pt",
                context=font_violations_sizes[0][0].text[:80],
            ))
        elif checked_paras > 0:
            result.passed += 1

        # ── 3. Межстрочный интервал ──────────────────────────────────
        required_spacing    = cfg.get("line_spacing", 1.5)
        spacing_tolerance   = cfg.get("line_spacing_tolerance", 0.05)

        spacing_violations = []
        for idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            if "heading" in style_name.lower() or "заголовок" in style_name.lower():
                continue

            sp = _get_para_spacing(para._element)
            if sp["multiplier"] is None:
                continue  # Интервал не задан явно — наследуется, пропускаем

            if abs(sp["multiplier"] - required_spacing) > spacing_tolerance:
                spacing_violations.append((idx, para, sp["multiplier"]))

        if spacing_violations:
            for idx, para, actual_sp in spacing_violations:
                result.violations.append(self._violation(
                    rule_id="GOST-LINE-SPACING",
                    severity=Severity.ERROR,
                    message=f"Межстрочный интервал {actual_sp}х — ожидается {required_spacing}х",
                    location=f"Параграф #{idx}: «{para.text[:50]}»",
                    suggestion=f"Установите полуторный межстрочный интервал ({required_spacing})",
                ))
        else:
            result.passed += 1

        # ── 4. Абзацный отступ (красная строка) ─────────────────────
        required_indent_mm  = cfg.get("first_line_indent_mm", 12.5)
        indent_tolerance_mm = cfg.get("first_line_indent_tolerance_mm", 1.5)

        # Строим карту element → ParsedParagraph для получения страницы/абзаца
        parsed_para_map = {pp.text.strip(): pp for pp in document.non_empty_paragraphs} if document else {}

        any_indent_issue = False
        for idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            if "heading" in style_name.lower() or "заголовок" in style_name.lower():
                continue

            indent_mm = _get_para_first_line_indent(para._element)
            if indent_mm is None:
                continue

            # Получаем ParsedParagraph для страницы/абзаца
            pp = parsed_para_map.get(para.text.strip())
            loc = pp.location_str if pp else f"Параграф #{idx}"
            preview = para.text.strip()[:60]
            page_num = pp.page_number if pp else None
            para_on_p = pp.para_on_page if pp else None

            if indent_mm == 0:
                any_indent_issue = True
                v = self._violation(
                    rule_id="GOST-INDENT-MISSING",
                    severity=Severity.WARNING,
                    message="Отсутствует абзацный отступ (красная строка)",
                    location=f"{loc}: «{preview}»",
                    suggestion=f"Установите отступ первой строки {required_indent_mm} мм (1.25 см)",
                    context=preview,
                )
                v.page_number = page_num
                v.para_on_page = para_on_p
                v.global_index = pp.index if pp else idx
                result.violations.append(v)
            elif abs(indent_mm - required_indent_mm) > indent_tolerance_mm:
                any_indent_issue = True
                v = self._violation(
                    rule_id="GOST-INDENT-WRONG",
                    severity=Severity.WARNING,
                    message=f"Неверный абзацный отступ: {indent_mm} мм (ожидается {required_indent_mm} мм)",
                    location=f"{loc}: «{preview}»",
                    suggestion=f"Установите отступ первой строки ровно {required_indent_mm} мм",
                    context=preview,
                )
                v.page_number = page_num
                v.para_on_page = para_on_p
                v.global_index = pp.index if pp else idx
                result.violations.append(v)

        if not any_indent_issue:
            result.passed += 1

        return result

    # ── Вспомогательный метод для проверки полей ─────────────────────
    def _check_margin(
        self,
        result: CheckResult,
        margins: dict,
        side: str,
        expected_mm: float,
        label: str,
    ) -> None:
        actual = margins.get(side)
        tolerance = self.config.get("margin_tolerance_mm", 1.5)

        if actual is None:
            result.violations.append(self._violation(
                rule_id=f"GOST-MARGIN-{side.upper()}",
                severity=Severity.WARNING,
                message=f"Не удалось прочитать {label} поле",
            ))
            return

        if abs(actual - expected_mm) > tolerance:
            result.violations.append(self._violation(
                rule_id=f"GOST-MARGIN-{side.upper()}",
                severity=Severity.ERROR,
                message=(
                    f"{label.capitalize()} поле: {actual} мм — "
                    f"ожидается {expected_mm} мм (ГОСТ 7.32-2017)"
                ),
                suggestion=f"Установите {label} поле {expected_mm} мм",
            ))
        else:
            result.passed += 1
