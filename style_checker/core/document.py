"""
core/document.py
Парсер DOCX-документа.
Поддерживает: параграфы, таблицы, изображения, номера страниц.
"""

from dataclasses import dataclass, field
from typing import Optional
import re
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from io import BytesIO


@dataclass
class Paragraph:
    index:       int            # Глобальный порядковый номер
    text:        str
    style_name:  str
    level:       Optional[int] = None
    font_size:   Optional[float] = None
    bold:        bool = False
    italic:      bool = False
    alignment:   Optional[str] = None
    list_level:  Optional[int] = None
    # Новые поля
    page_number: int = 1        # Номер страницы
    para_on_page: int = 1       # Порядковый номер параграфа на этой странице
    in_table:    bool = False   # Находится внутри таблицы
    table_index: Optional[int] = None  # Индекс таблицы (если in_table)
    is_image:    bool = False   # Параграф содержит изображение

    @property
    def is_heading(self) -> bool:
        return self.level is not None

    @property
    def is_empty(self) -> bool:
        return not self.text.strip()

    @property
    def location_str(self) -> str:
        """Читаемое описание места: страница, абзац на странице, текст."""
        loc = f"Стр. {self.page_number}, абз. {self.para_on_page}"
        if self.in_table:
            loc += f" [таблица {self.table_index + 1}]"
        if self.is_image:
            loc += " [изображение]"
        return loc

    @property
    def context_preview(self) -> str:
        """Первые 80 символов текста для показа в отчёте."""
        t = self.text.strip()
        return (t[:77] + "…") if len(t) > 80 else t


@dataclass
class ParsedDocument:
    paragraphs:  list[Paragraph] = field(default_factory=list)
    properties:  dict = field(default_factory=dict)
    page_count:  int = 1        # Оценочное количество страниц
    table_count: int = 0
    image_count: int = 0
    skip_pages:  set[int] = field(default_factory=set)  # Страницы для пропуска

    @property
    def headings(self) -> list[Paragraph]:
        return [p for p in self.paragraphs if p.is_heading and p.page_number not in self.skip_pages]

    @property
    def non_empty_paragraphs(self) -> list[Paragraph]:
        return [p for p in self.paragraphs
                if not p.is_empty and p.page_number not in self.skip_pages]

    def headings_by_level(self, level: int) -> list[Paragraph]:
        return [p for p in self.paragraphs
                if p.level == level and p.page_number not in self.skip_pages]

    def text_paragraphs(self) -> list[Paragraph]:
        return [p for p in self.paragraphs
                if not p.is_heading and not p.is_empty
                and p.page_number not in self.skip_pages]


# ── Helpers ──────────────────────────────────────────────────────────────────

def _get_heading_level(style_name: str) -> Optional[int]:
    for pat in [r"[Hh]eading\s*(\d)", r"Заголовок\s*(\d)"]:
        m = re.match(pat, style_name)
        if m:
            return int(m.group(1))
    return None


def _get_font_size(para) -> Optional[float]:
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    return None


def _get_alignment(para) -> Optional[str]:
    alignment_map = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY"}
    if para.alignment is not None:
        return alignment_map.get(para.alignment.value, str(para.alignment))
    return None


def _get_list_level(para) -> Optional[int]:
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    return int(ilvl.get(qn("w:val"), 0)) if ilvl is not None else 0


def _has_image(para_element) -> bool:
    """Проверяет наличие изображения в параграфе."""
    return (
        para_element.find(".//" + qn("w:drawing")) is not None or
        para_element.find(".//" + qn("w:pict")) is not None
    )


def _has_page_break(para_element) -> bool:
    """Параграф начинается с новой страницы."""
    # Явный разрыв страницы через w:br type=page
    for run in para_element.findall(qn("w:r")):
        for br in run.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    # pageBreakBefore в свойствах параграфа
    pPr = para_element.find(qn("w:pPr"))
    if pPr is not None:
        pbr = pPr.find(qn("w:pageBreakBefore"))
        if pbr is not None:
            val = pbr.get(qn("w:val"), "true")
            if val not in ("false", "0"):
                return True
    return False


def _make_paragraph(
    para, index: int, page: int, para_on_page: int,
    in_table: bool = False, table_index: Optional[int] = None
) -> Paragraph:
    try:
        style_name = para.style.name if para.style else "Normal"
    except Exception:
        style_name = "Normal"
    has_img = _has_image(para._element)
    return Paragraph(
        index=index,
        text=para.text,
        style_name=style_name,
        level=_get_heading_level(style_name),
        font_size=_get_font_size(para),
        bold=any(r.bold for r in para.runs),
        italic=any(r.italic for r in para.runs),
        alignment=_get_alignment(para),
        list_level=_get_list_level(para),
        page_number=page,
        para_on_page=para_on_page,
        in_table=in_table,
        table_index=table_index,
        is_image=has_img,
    )


# ── Parser ────────────────────────────────────────────────────────────────────

class DocumentParser:
    def parse(self, source, skip_pages: set[int] = None) -> ParsedDocument:
        if isinstance(source, bytes):
            source = BytesIO(source)

        doc = DocxDocument(source)
        parsed = ParsedDocument(skip_pages=skip_pages or set())

        core_props = doc.core_properties
        parsed.properties = {
            "author":   core_props.author or "",
            "title":    core_props.title or "",
            "created":  str(core_props.created or ""),
            "modified": str(core_props.modified or ""),
        }

        # Строим индекс таблиц: элемент XML → порядковый номер
        table_elements = doc.element.body.findall(".//" + qn("w:tbl"))
        table_index_map = {id(t._element): i for i, t in enumerate(doc.tables)}
        parsed.table_count = len(doc.tables)
        parsed.image_count = len(doc.inline_shapes)

        page = 1
        para_on_page = 0
        global_index = 0

        # Карта element → docx Paragraph (у них корректный .part/.style)
        para_el_map = {p._element: p for p in doc.paragraphs}
        # Карта element → индекс таблицы
        tbl_el_map  = {t._element: i for i, t in enumerate(doc.tables)}

        for body_child in doc.element.body:
            tag = body_child.tag.split("}")[-1] if "}" in body_child.tag else body_child.tag

            if tag == "p":
                para = para_el_map.get(body_child)
                if para is None:
                    continue

                if _has_page_break(body_child):
                    page += 1
                    para_on_page = 0

                para_on_page += 1
                p = _make_paragraph(para, global_index, page, para_on_page)
                parsed.paragraphs.append(p)
                global_index += 1

            elif tag == "tbl":
                tbl_order = tbl_el_map.get(body_child)
                if tbl_order is None:
                    continue
                tbl = doc.tables[tbl_order]

                seen_cells = set()
                for row in tbl.rows:
                    for cell in row.cells:
                        if id(cell) in seen_cells:
                            continue
                        seen_cells.add(id(cell))
                        for cell_para in cell.paragraphs:
                            if _has_page_break(cell_para._element):
                                page += 1
                                para_on_page = 0
                            para_on_page += 1
                            p = _make_paragraph(
                                cell_para, global_index, page, para_on_page,
                                in_table=True, table_index=tbl_order
                            )
                            parsed.paragraphs.append(p)
                            global_index += 1

        parsed.page_count = page
        return parsed
